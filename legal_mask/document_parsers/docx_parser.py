from __future__ import annotations
from pathlib import Path
import uuid
from docx import Document
from legal_mask.document_parsers.base import BaseParser
from legal_mask.types import DocumentModel


class DocxParser(BaseParser):
    def parse(self, path: Path) -> DocumentModel:
        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs]
        content = "\n".join(paragraphs)
        return DocumentModel(
            id=str(uuid.uuid4()),
            filename=path.name,
            content=content,
            original_path=str(path),
            file_type="docx",
        )

    def apply_masked_content(self, path: Path, masked_text: str, output_path: Path):
        doc = Document(str(path))
        lines = masked_text.split("\n")
        for i, paragraph in enumerate(doc.paragraphs):
            if i < len(lines):
                paragraph.text = lines[i]
        doc.save(str(output_path))
