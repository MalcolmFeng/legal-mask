from __future__ import annotations
from pathlib import Path
from legal_mask.types import DocumentModel, Annotation, MaskConfig
from legal_mask.engine.masker import Masker
from docx import Document


class Exporter:
    @staticmethod
    def export_text(doc: DocumentModel, annotations: list[Annotation], config: MaskConfig | None = None) -> str:
        return Masker.apply(doc.content, annotations, config)

    @staticmethod
    def export_docx(doc: DocumentModel, annotations: list[Annotation], output_path: str | Path, config: MaskConfig | None = None):
        masked_text = Masker.apply(doc.content, annotations, config)
        if doc.file_type == "docx":
            source = Document(doc.original_path)
            lines = masked_text.split("\n")
            for i, paragraph in enumerate(source.paragraphs):
                if i < len(lines):
                    paragraph.text = lines[i]
            source.save(str(output_path))
        else:
            out = Document()
            for line in masked_text.split("\n"):
                out.add_paragraph(line)
            out.save(str(output_path))

    @staticmethod
    def export_comparison(doc: DocumentModel, annotations: list[Annotation], config: MaskConfig | None = None) -> str:
        masked = Masker.apply(doc.content, annotations, config)
        lines = []
        lines.append("=" * 60)
        lines.append("原文:")
        lines.append("=" * 60)
        lines.append(doc.content)
        lines.append("")
        lines.append("=" * 60)
        lines.append("脱敏后:")
        lines.append("=" * 60)
        lines.append(masked)
        return "\n".join(lines)
