from legal_mask.engine.exporter import Exporter
from legal_mask.engine.masker import Masker
from legal_mask.types import DocumentModel, Annotation, AnnotationStatus, SensitiveType
import uuid
from pathlib import Path


def make_annotation(stype, start, end, text, status=AnnotationStatus.CONFIRMED):
    return Annotation(
        id=str(uuid.uuid4()), doc_id="test",
        sensitive_type=stype, start=start, end=end,
        text=text, confidence=1.0, source="test", status=status,
    )


def test_export_text():
    doc = DocumentModel(id="1", filename="test.txt", content="原告张三", original_path="", file_type="txt")
    anns = [make_annotation(SensitiveType.PERSON_NAME, 2, 4, "张三")]
    result = Exporter.export_text(doc, anns)
    assert result == "原告[姓名]"


def test_export_docx(tmp_path):
    from docx import Document
    source = tmp_path / "source.docx"
    doc_obj = Document()
    doc_obj.add_paragraph("原告张三")
    doc_obj.save(str(source))

    doc = DocumentModel(id="1", filename="source.docx", content="原告张三", original_path=str(source), file_type="docx")
    anns = [make_annotation(SensitiveType.PERSON_NAME, 2, 4, "张三")]
    output = tmp_path / "masked.docx"
    Exporter.export_docx(doc, anns, str(output))

    result_doc = Document(str(output))
    assert result_doc.paragraphs[0].text == "原告[姓名]"
